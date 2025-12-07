from docx import Document
from openpyxl import load_workbook
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import os

file_path = os.path.join(os.path.dirname(__file__), '수료증명단.xlsx')
load_wb = load_workbook(file_path)
load_ws = load_wb.active

name_list = []
birthday_list = []
ho_list = []
for i in range(1, load_ws.max_row + 1):
    name_list.append(load_ws.cell(i, 1).value)
    birthday_list.append(load_ws.cell(i, 2).value)
    ho_list.append(load_ws.cell(i, 3).value)

print(name_list)
print(birthday_list)
print(ho_list)

for i in range(len(name_list)):
    file_path = os.path.join(os.path.dirname(__file__), '수료증양식.docx')
    doc = docx.Document(file_path)

    style = doc.styles['Normal']
    style.font.name = '나눔고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(12)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('       제 ' + ho_list[i] + ' 호\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(14)

    para = doc.add_paragraph()
    run = para.add_run('\n')
    run = para.add_run('수  료  증\n')
    run.font.name = '나눔고딕'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(30)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n')
    run = para.add_run('     성    명: ' + name_list[i] + '\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(14)
    run = para.add_run('     생년월일: ' + birthday_list[i] + '\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(14)
    run = para.add_run('     교육과정: 메가스터디 파이썬 응용프로젝트반\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(14)
    run = para.add_run('     교육날짜: 2026.01.05~2026.01.30\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(14)

    para = doc.add_paragraph()
    run = para.add_run('\n')
    run = para.add_run('     위 사람은 메가스터디 파이썬 과정을 이수\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('   하였으므로 이 증서를 수여합니다.\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n')
    run = para.add_run('2026년 1월 30일')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n\n')
    run = para.add_run('메가스터디 파이썬반 선생님')
    run.font.name = '나눔고딕'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(os.path.join(os.path.dirname(__file__), name_list[i] + '.docx'))
    convert(os.path.join(os.path.dirname(__file__), name_list[i] + '.docx'),
            os.path.join(os.path.dirname(__file__), name_list[i] + '.pdf'))
